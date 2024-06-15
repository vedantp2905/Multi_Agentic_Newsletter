import os
import docx
from langchain_openai import OpenAI
import streamlit as st  # Import Streamlit for web application interface
from docx import Document  # Import python-docx for Word document creation
from io import BytesIO  # Import BytesIO for in-memory file operations
import replicate  # Import Replicate for image generation
import requests  # Import requests to download images

import asyncio
import google.generativeai as genai  # Import the appropriate module for Gemini

# Import ChatGoogleGenerativeAI: High-level interface to Google's AI models
from langchain_google_genai import ChatGoogleGenerativeAI

# Import Agent, Task, Crew, Process: Core classes for AI agent workflows
from crewai import Agent, Task, Crew, Process

# Import DuckDuckGoSearchRun: Tool for web searches via DuckDuckGo
from langchain_community.tools import DuckDuckGoSearchRun

# Function to generate text based on topic
def generate_text(llm, topic):
    inputs = {'topic': topic}

    # Initialize DuckDuckGo web search tool: Enables real-time fact-finding for debates
    search_tool = DuckDuckGoSearchRun(
        name="duckduckgo_search",
        description="""Search the web using DuckDuckGo. Action Input should look like this:
                       {"query": "<Whatever you want to search>"}"""
    )

    # Define Blog Researcher Agent
    blog_researcher = Agent(
        role='Blog Content Researcher',
        goal='Conduct thorough research to uncover compelling insights for engaging blog content.',
        backstory=("An experienced content strategist with a knack for analyzing trends and audience behavior, "
                   "delivering actionable insights for high-quality blog content."),
        verbose=True,
        allow_delegation=False,
        llm=llm
    )

    # Define Blog Writer Agent
    blog_writer = Agent(
        role='Blog Writer',
        goal='Craft authoritative and engaging blog content that resonates with the audience and establishes the brand as a leader.',
        backstory=("A seasoned writer known for distilling complex topics into captivating stories, with a deep understanding of audience psychology."),
        verbose=True,
        allow_delegation=False,
        llm=llm
    )

    # Define Blog Reviewer Agent
    blog_reviewer = Agent(
        role='Content Reviewer',
        goal='Review and refine blog drafts to ensure they meet high standards of quality and impact.',
        backstory=("An expert editor with a meticulous eye for detail, known for elevating content to publication-ready standards."),
        verbose=True,
        allow_delegation=False,
        llm=llm
    )

    # Define Task for Researcher
    task_researcher = Task(
        description=(f"Research the latest trends and insights on {topic}. Identify key developments, emerging trends, unique perspectives, and content ideas."),
        agent=blog_researcher,
        expected_output=(
            f"1. Overview and background of {topic}.\n"
            "2. Recent key developments.\n"
            "3. Emerging trends and innovative approaches.\n"
            "4. Unique angles and untapped opportunities.\n"
            "5. Potential content ideas with brief descriptions.\n"
            "6. List of relevant sources."
        ),
        tools=[search_tool]
    )

    # Define Task for Writer
    task_writer = Task(
        description=(f"Based on the research report, craft an engaging and authoritative blog post on {topic}."),
        agent=blog_writer,
        expected_output=(
            "1. Engaging introduction with a hook.\n"
            "2. Detailed exploration of key developments.\n"
            "3. Use of emerging trends and innovative ideas in content.\n"
            "4. Use of unique angles and perspectives in content.\n"
            "5. Clear explanations of complex concepts.\n"
            "7. Compelling conclusion.\n"
        )
    )

    # Define Task for Reviewer
    task_reviewer = Task(
        description=(f"Review the drafted blog content on {topic}, providing detailed feedback and revisions for quality and impact."),
        agent=blog_reviewer,
        expected_output=(
            "1. Overall content assessment.\n"
            "2. Identification of inaccuracies and gaps.\n"
            "3. Suggestions for improving flow and readability.\n"
            "4. Recommendations for tone and voice.\n"
            "5. Edits for grammar and punctuation.\n"
            "6. Feedback on multimedia use.\n"
            "7. Final assessment of readiness."
        )
    )

    # Define Task for Final Writer
    task_final_writer = Task(
        description=(f"Revise the blog content on {topic} based on the reviewer's feedback, ensuring it meets high standards."),
        agent=blog_writer,
        expected_output=(
            "1. Factually accurate and corrected content.\n"
            "2. Clear, well-structured flow.\n"
            "3. Concise and engaging language.\n"
            "4. Consistent tone and voice.\n"
            "5. Enhanced insights.\n"
            "6. Addressed reviewer feedback.\n"
            "7. Creative and engaging blog title.\n"
            "8. Final draft of at least 1000 words."
        )
    )

    # Initialize Crew: Coordinates agents and tasks for structured blog content workflow
    crew = Crew(
        agents=[blog_researcher, blog_writer, blog_reviewer, blog_writer],
        tasks=[task_researcher, task_writer, task_reviewer, task_final_writer],
        verbose=2,
        context={"Blog Topic is ": topic}
    )

    # Start the workflow and generate the result
    result = crew.kickoff(inputs=inputs)

    return result

# Function to generate images based on prompts
def generate_images(replicate_api_token, prompt):
    
    os.environ["REPLICATE_API_TOKEN"] = replicate_api_token

    # Define the input for the image generation
    input = {
        "prompt": prompt,
        "scheduler": "K_EULER"
    }

    # Generate the image
    output = replicate.run(
        "stability-ai/stable-diffusion:ac732df83cea7fff18b8472768c88ad041fa750ff7682a21affe81863cbe77e4",
        input=input
    )

    # Assuming output is a list of URLs, return the first one
    if output and isinstance(output, list) and len(output) > 0:
        return output[0]
    else:
        raise ValueError("No image URL returned from Replicate API.")

# Streamlit web application
def main():
   st.header('AI Blog Content Generator')
   mod = None
   with st.sidebar:
       with st.form('Gemini/OpenAI'):
            # User selects the model (Gemini/Cohere) and enters API keys
            model = st.radio('Choose Your LLM', ('Gemini', 'OpenAI'))
            api_key = st.text_input(f'Enter your API key', type="password")
            replicate_api_token = st.text_input('Enter Replicate API key', type="password")
            submitted = st.form_submit_button("Submit")

   # Check if API key is provided and set up the language model accordingly
   if api_key:
        if model == 'OpenAI':
            async def setup_OpenAI():
                loop = asyncio.get_event_loop()
                if loop is None:
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)

                os.environ["OPENAI_API_KEY"] = api_key
                llm = OpenAI(temperature=0.6,max_tokens=3500)
                return llm

            llm = asyncio.run(setup_OpenAI())
            mod = 'OpenAI'

        elif model == 'Gemini':
            async def setup_gemini():
                loop = asyncio.get_event_loop()
                if loop is None:
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)

                llm = ChatGoogleGenerativeAI(
                    model="gemini-1.5-flash",
                    verbose=True,
                    temperature=0.6,
                    google_api_key=api_key  # Use the API key from the environment variable
                )
                return llm

            llm = asyncio.run(setup_gemini())
            mod = 'Gemini'
        
        # User input for the blog topic
        topic = st.text_input("Enter the blog topic:")

        if st.button("Generate Blog Content"):
            with st.spinner("Generating content..."):
                generated_content = generate_text(llm, topic)
                generated_image_url = generate_images(replicate_api_token, topic)

                content_lines = generated_content.split('\n')
                first_line = content_lines[0]
                remaining_content = '\n'.join(content_lines[1:])

                st.markdown(first_line)
                st.image(generated_image_url, caption="Generated Image", use_column_width=True)
                st.markdown(remaining_content)

                # Download the images and add them to the document
                response = requests.get(generated_image_url)
                image = BytesIO(response.content)

                doc = Document()

                # Option to download content as a Word document
                doc.add_heading(topic, 0)
                doc.add_paragraph(first_line)
                doc.add_picture(image, width=docx.shared.Inches(6))  # Add image to the document
                doc.add_paragraph(remaining_content)

                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                st.download_button(
                    label="Download as Word Document",
                    data=buffer,
                    file_name=f"{topic}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

if __name__ == "__main__":
    main()