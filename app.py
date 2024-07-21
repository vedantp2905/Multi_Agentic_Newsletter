import os
import asyncio
from langchain_groq import ChatGroq
import requests
from typing import Type, Any
from io import BytesIO
from crewai_tools import ScrapeWebsiteTool
import streamlit as st
from streamlit import cache
from docx import Document
from pydantic.v1 import BaseModel, Field
from crewai_tools.tools.base_tool import BaseTool
from langchain_openai import ChatOpenAI
from langchain_google_genai import ChatGoogleGenerativeAI
from crewai import Agent, Task, Crew

def verify_gemini_api_key(api_key):
    API_VERSION = 'v1'
    api_url = f"https://generativelanguage.googleapis.com/{API_VERSION}/models?key={api_key}"
    
    try:
        response = requests.get(api_url, headers={'Content-Type': 'application/json'})
        response.raise_for_status()  # Raises an HTTPError for bad responses
        
        # If we get here, it means the request was successful
        return True
    
    except requests.exceptions.HTTPError as e:
        
        return False
    
    except requests.exceptions.RequestException as e:
        # For any other request-related exceptions
        raise ValueError(f"An error occurred: {str(e)}")

def verify_gpt_api_key(api_key):
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    
    # Using a simple request to the models endpoint
    response = requests.get("https://api.openai.com/v1/models", headers=headers)
    
    if response.status_code == 200:
        return True
    elif response.status_code == 401:
        return False
    else:
        print(f"Unexpected status code: {response.status_code}")
        return False
    
def verify_groq_api_key(api_key):
    api_url = "https://api.groq.com/openai/v1/models"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    
    try:
        response = requests.get(api_url, headers=headers)
        response.raise_for_status()  # Raises an HTTPError for bad responses
        
        # If we get here, it means the request was successful
        return True
    
    except requests.exceptions.HTTPError as e:
        
        return False
    
    except requests.exceptions.RequestException as e:
        # For any other request-related exceptions
        raise ValueError(f"An error occurred: {str(e)}")

import requests

def verify_serpapi_key(api_key):
    
    base_url = "https://serpapi.com/search"
    
    # Parameters for a simple Google search
    params = {
        "engine": "google",
        "q": "test",
        "api_key": api_key
    }
    
    try:
        response = requests.get(base_url, params=params)
        
        if response.status_code == 200:
            # If we get a 200 OK response, the API key is valid
            return True
        elif response.status_code == 401:
            # 401 Unauthorized typically means invalid API key
            return False
        else:
            # For other status codes, we'll assume the key is invalid
            print(f"Unexpected status code: {response.status_code}")
            return False
    
    except requests.exceptions.RequestException as e:
        # Handle any request exceptions
        print(f"An error occurred: {str(e)}")
        return False


serp_api_key = ''

class SerpApiGoogleSearchToolSchema(BaseModel):
    q: str = Field(..., description="Parameter defines the query you want to search. You can use anything that you would use in a regular Google search. e.g. inurl:, site:, intitle:.")
    tbs: str = Field("qdr:w2", description="Time filter to limit the search to the last two weeks.")

class SerpApiGoogleSearchTool(BaseTool):
    name: str = "Google Search"
    description: str = "Search the internet"
    args_schema: Type[BaseModel] = SerpApiGoogleSearchToolSchema
    search_url: str = "https://serpapi.com/search"
    
    def _run(
        self,
        q: str,
        tbs: str = "qdr:w2",
        **kwargs: Any,
    ) -> Any:
        global serp_api_key
        payload = {
            "engine": "google",
            "q": q,
            "tbs": tbs,
            "api_key": serp_api_key,
        }
        headers = {
            'content-type': 'application/json'
        }
    
        response = requests.get(self.search_url, headers=headers, params=payload)
        results = response.json()
    
        summary = ""
        for key in ['answer_box_list', 'answer_box', 'organic_results', 'sports_results', 'knowledge_graph', 'top_stories']:
            if key in results:
                summary += str(results[key])
                break
        
        print(summary)
        
        return summary
   
def generate_text(llm, topic, serpapi_key):
    
    inputs = {'topic': topic}
    
    search_tool = SerpApiGoogleSearchTool()
    
    scrape_tool = ScrapeWebsiteTool()

    researcher_agent = Agent(
        role='Newsletter Content Researcher',
        goal='Search for 5 stories on the given topic, find unique 5 URLs containing the stories, and scrape relevant information from these URLs.',
        backstory=(
            "An experienced researcher with strong skills in web scraping, fact-finding, and "
            "analyzing recent trends to provide up-to-date information for high-quality newsletters."
        ),
        verbose=True,
        allow_delegation=False,
        max_iter = 5,
        llm=llm
    )

    writer_agent = Agent(
        role='Content Writer',
        goal='Write detailed, engaging, and informative summaries of the stories found by the researcher using the format specified.',
        backstory=("An experienced writer with a background in journalism and content creation. "
                   "Skilled in crafting compelling narratives and distilling complex information into "
                   "accessible formats. Adept at conducting research and synthesizing insights for engaging content."),
        verbose=True,
        allow_delegation=False,
        max_iter = 5,
        llm=llm
    )

    reviewer_agent = Agent(
        role='Content Reviewer',
        goal='Review and refine content drafts to ensure they meet high standards of quality and impact like major newsletters.',
        backstory=("A meticulous reviewer with extensive experience in editing and proofreading, "
                   "known for their keen eye for detail and commitment to maintaining the highest quality standards in published content."),
        verbose=True,
        allow_delegation=False,
        max_iter = 5,
        llm=llm
    )
    
    final_writer_agent = Agent(
        role='Final Content Writer',
        goal='Compile, refine, and structure all reviewed and approved content into a cohesive and engaging newsletter format. Ensure that the final product is polished, logically structured, and ready for publication, providing a seamless and informative reading experience for the audience.',
        backstory=("An accomplished writer and editor with extensive experience in journalism, content creation, and editorial management. "
                   "Known for their ability to craft compelling narratives and ensure consistency and quality across all sections of a publication. "
                   "With a keen eye for detail and a deep understanding of audience engagement, this writer excels in transforming raw content into polished, professional-grade newsletters that captivate readers and deliver clear, valuable insights."),
        verbose=True,
        allow_delegation=False,
        llm=llm,
        max_iter = 5
    )
    
    task_researcher = Task(
        description=(f'Research and identify the most interesting 5 stories on the topic of {topic} '
                     'Scrape detailed content from relevant websites to gather comprehensive material.'
                     'If unable to scrape a URL, find a new URL and scrape it.'),
        agent=researcher_agent,
        expected_output=('A list of recent 5 stories within last 2 weeks with their respective website URLs. '
                         'Scraped content from all URLs that can be used further by the writer.'),
        tools=[search_tool, scrape_tool]
    )

    task_writer = Task(
        description=('Write detailed summaries of the stories found by the researcher. '
                     'Ensure each summary is informative, engaging, and provides clear insights into the story.'),
        agent=writer_agent,
        expected_output=('Summarized content for all the stories, each summary being 150-200 words long, '
                         'with clear and concise information.'
                         'Links to original source found by researcher and any additional information if needed.')
    )

    task_reviewer = Task(
        description=('Review the summarized content provided by the writer for accuracy, coherence, and quality. '
                     'Ensure that the content is free from errors and meets the publication standards.'),
        agent=reviewer_agent,
        expected_output=('Reviewed content with suggestions for improvements, if any. '
                         'Final versions of summaries that are ready for inclusion in the newsletter.'
                         'Verify the links are opening to correct pages')
    )

    task_final_writer = Task(
        description=('Compile the reviewed and refined content into a well-structured newsletter format. '
                     'Ensure the newsletter is visually appealing and flows logically from one section to the next.'),
        agent=final_writer_agent,
        expected_output=(
            """Final newsletter document with all the reviewed summaries, formatted and ready for publication. 
            The newsletter should include:
            - Introduction: A compelling hook sentence to engage readers and make the newsletter fun.
            - Contents section:
                - Summarize each story in one interesting sentence.
            - Main content sections (1 highlight story and 4 other stories):
                -Highlight one of the stories as the most interesitng one in the newsletter which you feel desevres to be the highlight. Explicitly mention it as newsletters highlight at the start of the main conetent section
                - Each story should have:
                    - A small introduction.
                    - Details presented in 3-4 bullet points.
                    - Explanation of why it matters or a call to action
                    - Links to original source found by researcher and any additional information if needed.
            - Conclusion:
                - Wrap up the newsletter by summarizing all content and providing a final thought or conclusion.
            """
        )
    )


    crew = Crew(
        agents=[researcher_agent, writer_agent, reviewer_agent, final_writer_agent],
        tasks=[task_researcher, task_writer, task_reviewer, task_final_writer],
        verbose=2,
        context={"Newsletter Topic is ": topic}
    )

    result = crew.kickoff(inputs=inputs)

    return result

def main():
    
    st.header('AI Newsletter Content Generator')

    global serp_api_key
    validity_serpapi = False
    validity_model = False
    
    # Initialize session state
    if 'generated_content' not in st.session_state:
        st.session_state.generated_content = None
    if 'topic' not in st.session_state:
        st.session_state.topic = ""
    
    with st.sidebar:
        with st.form('Gemini/OpenAI/Groq'):
            model = st.radio('Choose Your LLM', ('Gemini', 'OpenAI','Groq'))
            api_key = st.text_input(f'Enter your API key', type="password")
            serp_api_key = st.text_input(f'Enter your SerpAPI key', type="password")
            submitted = st.form_submit_button("Submit")

        if api_key and serp_api_key:
            if model == "Gemini":
                validity_model = verify_gemini_api_key(api_key)
                if validity_model ==True:
                    st.write(f"Valid {model} API key")
                else:
                    st.write(f"Invalid {model} API key")
            elif model == "OpenAI":
                validity_model = verify_gpt_api_key(api_key)
                if validity_model ==True:
                    st.write(f"Valid {model} API key")
                else:
                    st.write(f"Invalid {model} API key")            
            elif model == "Groq":
                validity_model = verify_groq_api_key(api_key)
                if validity_model ==True:
                    st.write(f"Valid {model} API key")
                else:
                    st.write(f"Invalid {model} API key")
            
            validity_serpapi = verify_serpapi_key(serp_api_key)
            if validity_serpapi ==True:
                st.write(f"Valid SerpAPI API key")
            else:
                st.write(f"Invalid SerpAPI API key")        
        
    if validity_model and validity_serpapi:
        if model == 'OpenAI':
            async def setup_OpenAI():
                loop = asyncio.get_event_loop()
                if loop is None:
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)

                os.environ["OPENAI_API_KEY"] = api_key
                llm = ChatOpenAI(model='gpt-4-turbo',temperature=0.6, max_tokens=2000,api_key=api_key)
                print("OpenAI Configured")
                return llm

            llm = asyncio.run(setup_OpenAI())

        elif model == 'Gemini':
            async def setup_gemini():
                loop = asyncio.get_event_loop()
                if loop is None:
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)

                llm = ChatGoogleGenerativeAI(
                    model="gemini-1.5-flash",
                    verbose=True,
                    temperature=0.6,
                    google_api_key=api_key
                )
                print("Gemini Configured")
                return llm

            llm = asyncio.run(setup_gemini())
            
        elif model == 'Groq':
            async def setup_groq():
                loop = asyncio.get_event_loop()
                if loop is None:
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)

                llm = ChatGroq(
                    api_key = api_key,
                    model = 'llama3-70b-8192'
                )
                return llm

            llm = asyncio.run(setup_groq())
            
        topic = st.text_input("Enter the newsletter topic:")
        st.session_state.topic = topic

        if st.button("Generate Newsletter Content"):
            with st.spinner("Generating content..."):
                st.session_state.generated_content = generate_text(llm, topic,serp_api_key)

        if st.session_state.generated_content:
            
            st.markdown(st.session_state.generated_content)

            doc = Document()

            doc.add_heading(topic, 0)
            doc.add_paragraph(st.session_state.generated_content)
            

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)


            st.download_button(
    label="Download as Word Document",
    data=buffer,
    file_name=f"{topic}.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

if __name__ == "__main__":
    main()
